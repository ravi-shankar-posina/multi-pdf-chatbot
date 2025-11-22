"""
Test Script Pipeline - Flask API Module
BRD -> Summary -> Knowledge Graph -> Test Scripts
"""

import os
import re
import shutil
import pdfplumber
from docx import Document
from openai import OpenAI
from openpyxl.styles import PatternFill, Font, Alignment
from openpyxl import Workbook
from langchain.document_loaders import TextLoader
from langchain.text_splitter import TokenTextSplitter
from langchain_openai import ChatOpenAI
from langchain_experimental.graph_transformers import LLMGraphTransformer
from dotenv import load_dotenv

load_dotenv()

api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    raise Exception("ERROR: OPENAI_API_KEY not set!")

client = OpenAI(api_key=api_key)

# Output folders
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
EXTRACTED_FOLDER = os.path.join(BASE_DIR, "output_tests", "extracted")
SUMMARY_FOLDER = os.path.join(BASE_DIR, "output_tests", "summary")
KG_FOLDER = os.path.join(BASE_DIR, "output_tests", "kg")
OUTPUT_TXT_FOLDER = os.path.join(BASE_DIR, "output_tests", "tests_txt")
OUTPUT_XLSX_FOLDER = os.path.join(BASE_DIR, "output_tests", "tests_xlsx")

for p in [EXTRACTED_FOLDER, SUMMARY_FOLDER, KG_FOLDER, OUTPUT_TXT_FOLDER, OUTPUT_XLSX_FOLDER]:
    os.makedirs(p, exist_ok=True)


def clear_output_folders():
    folders = [EXTRACTED_FOLDER, SUMMARY_FOLDER, KG_FOLDER, OUTPUT_TXT_FOLDER, OUTPUT_XLSX_FOLDER]
    for folder in folders:
        for filename in os.listdir(folder):
            file_path = os.path.join(folder, filename)
            try:
                if os.path.isfile(file_path) or os.path.islink(file_path):
                    os.remove(file_path)
                elif os.path.isdir(file_path):
                    shutil.rmtree(file_path)
            except:
                pass


def extract_pdf_content(pdf_path):
    text = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                t = page.extract_text()
                if t:
                    text += f"\n--- Page {i} Text ---\n{t}\n"
                for t_i, table in enumerate(page.extract_tables(), start=1):
                    text += f"\n--- Table {t_i} ---\n"
                    for row in table:
                        text += "\t".join([c if c else "" for c in row]) + "\n"
    except Exception as e:
        raise Exception(f"PDF extraction error: {e}")
    return text


def extract_docx_content(docx_path):
    text = ""
    try:
        doc = Document(docx_path)
        for para in doc.paragraphs:
            if para.text and para.text.strip():
                text += para.text + "\n"
        for t_i, table in enumerate(doc.tables, start=1):
            text += f"\n--- Table {t_i} ---\n"
            for row in table.rows:
                text += "\t".join([c.text.strip() for c in row.cells]) + "\n"
    except Exception as e:
        raise Exception(f"DOCX extraction error: {e}")
    return text


def process_file_to_txt(file_path):
    base, ext = os.path.splitext(os.path.basename(file_path))
    ext = ext.lower()
    if ext == ".pdf":
        content = extract_pdf_content(file_path)
    elif ext == ".docx":
        content = extract_docx_content(file_path)
    else:
        raise ValueError("Only PDF/DOCX supported")
    
    txt_path = os.path.join(EXTRACTED_FOLDER, f"{base}.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(content)
    return txt_path, content


SUMMARY_PROMPT = """
You are an expert business analyst specializing in BRDs, PDDs, FDDs, SOPs.

Provide:
==========================
PART 1 — FULL DOCUMENT SUMMARY
==========================
- Section summaries, Business purpose, Key requirements
- Functional logic, Actors, roles, responsibilities
- System interactions (SAP, VEEVA, S/4HANA, etc.)
- Assumptions and dependencies

==========================
PART 2 — TABLE SUMMARIES
==========================
For EACH table, write 10–12 sentence process explanation.
DO NOT recreate the table.

INPUT — NON-TABLE TEXT:
{clean_text}

INPUT — TABLES:
{tables_block}
"""


def summarize_cleaned_text(txt_path):
    with open(txt_path, "r", encoding="utf-8") as f:
        text = f.read()

    table_pattern = r"---\s*Table\s*(\d+)\s*---(.*?)(?=---\s*Table|\Z)"
    tables = re.findall(table_pattern, text, re.DOTALL)
    clean_text = re.sub(r"---\s*Table\s*\d+.*?(?=---\s*Table|\Z)", "", text, flags=re.DOTALL).strip()

    tables_block = ""
    for num, content in tables:
        tables_block += f"\n\n### TABLE {num}\n{content.strip()}"

    prompt = SUMMARY_PROMPT.format(clean_text=clean_text, tables_block=tables_block)

    response = client.chat.completions.create(
        model="gpt-4.1",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.0,
    )

    summary = response.choices[0].message.content
    base = os.path.splitext(os.path.basename(txt_path))[0]
    summary_path = os.path.join(SUMMARY_FOLDER, f"{base}_summary.txt")
    
    with open(summary_path, "w", encoding="utf-8") as f:
        f.write(summary)
    return summary_path, summary


def generate_knowledge_graph(txt_path):
    loader = TextLoader(txt_path, encoding="utf-8")
    document = loader.load()[0]
    
    splitter = TokenTextSplitter(chunk_size=800, chunk_overlap=200)
    chunks = splitter.split_documents([document])
    
    llm = ChatOpenAI(temperature=0, model_name="gpt-4o-mini")
    transformer = LLMGraphTransformer(llm=llm)
    graph_docs = transformer.convert_to_graph_documents(chunks)

    base = os.path.splitext(os.path.basename(txt_path))[0]
    out = os.path.join(KG_FOLDER, f"{base}_KG.txt")

    kg_content = f"=== KNOWLEDGE GRAPH FOR {base} ===\n\n"
    total = 0
    for idx, g in enumerate(graph_docs):
        kg_content += f"\n######## GRAPH DOCUMENT {idx+1} ########\n--- RELATIONSHIPS ---\n"
        for rel in getattr(g, "relationships", []):
            try:
                kg_content += f"{rel.source.id} --{str(rel.type).upper()}--> {rel.target.id}\n"
                total += 1
            except:
                pass
    kg_content += f"\nTOTAL RELATIONSHIPS: {total}\n"

    with open(out, "w", encoding="utf-8") as f:
        f.write(kg_content)
    return out, kg_content


def generate_test_scripts(kg_path, summary_path):
    with open(kg_path, "r", encoding="utf-8") as f:
        graph_text = f.read()
    with open(summary_path, "r", encoding="utf-8") as f:
        brd_text = f.read()

    prompt = f"""
You are an expert QA Analyst specialized in SAP S/4HANA.

Generate EXACTLY 8–12 structured SAP test cases:

### OUTPUT STRUCTURE ###
Test Name    <Short Title>
Test Type    MANUAL
Test Author  

Steps:
Step | Description | Expected Result | Selection Screen Details 
Step 1 | <QA verb> | <Expected result> | <fields> 

### RULES ###
- Use QA verbs: Validate / Check / Verify / Ensure / Navigate / Enter
- Use SAP terms: Fiori App IDs, T-codes, Material, Plant, Order
- NO explanations. ONLY formatted test cases.

GRAPH NODES:
{graph_text}

SUMMARY BRD:
{brd_text}
"""

    response = client.chat.completions.create(
        model="gpt-4.1",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.1
    )

    output_text = response.choices[0].message.content
    base = os.path.splitext(os.path.basename(kg_path))[0].replace("_KG", "").upper()
    
    txt_out = os.path.join(OUTPUT_TXT_FOLDER, f"{base}_Tests.txt")
    with open(txt_out, "w", encoding="utf-8") as f:
        f.write(output_text)

    # Generate Excel
    xlsx_out = generate_test_excel(output_text, base)
    
    return txt_out, xlsx_out, output_text


def generate_test_excel(output_text, base):
    test_cases = [tc.strip() for tc in re.split(r"(?=Test Name)", output_text) if tc.strip()]

    wb = Workbook()
    ws = wb.active
    ws.title = "Test Cases"

    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True)

    def style(cell):
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")

    ws.append(["SI No", "Test Case Name"])
    style(ws["A1"]); style(ws["B1"])

    for i, tc in enumerate(test_cases, start=1):
        m = re.search(r"Test Name\s*(.*)", tc)
        tname = m.group(1).strip() if m else f"Test Case {i}"
        ws.append([f"Test Case {i}", tname])

    for i, tc in enumerate(test_cases, start=1):
        sheet = wb.create_sheet(f"TestScript{i}")
        lines = tc.split("\n")
        row = 1

        for h in ["Test Name", "Test Type", "Test Author"]:
            for line in lines:
                if line.startswith(h):
                    sheet.cell(row=row, column=1).value = h
                    sheet.cell(row=row, column=2).value = line.replace(h, "").strip()
                    style(sheet.cell(row=row, column=1))
                    row += 1
                    break

        sheet.append(["Step", "Description", "Expected Result", "Selection Screen Details"])
        for col in range(1, 5):
            style(sheet.cell(row=row, column=col))
        row += 1

        for idx, step in enumerate(re.findall(r"Step\s*\d+\s*\|\s*([^|]+)\|\s*([^|]+)\|\s*(.*)", tc), start=1):
            sheet.append([f"Step {idx}"] + [s.strip() for s in step])

    for sheet in wb.worksheets:
        sheet.column_dimensions['A'].width = 20
        sheet.column_dimensions['B'].width = 70
        sheet.column_dimensions['C'].width = 50
        sheet.column_dimensions['D'].width = 40

    xlsx_out = os.path.join(OUTPUT_XLSX_FOLDER, f"{base}_Tests.xlsx")
    wb.save(xlsx_out)
    return xlsx_out


def run_test_script_pipeline(file_path):
    """Main pipeline function"""
    result = {
        "success": False,
        "extracted_path": None,
        "extracted_content": None,
        "summary_path": None,
        "summary_content": None,
        "kg_path": None,
        "kg_content": None,
        "test_txt_path": None,
        "test_xlsx_path": None,
        "test_scripts": None,
        "error": None
    }

    try:
        clear_output_folders()

        txt_path, extracted = process_file_to_txt(file_path)
        result["extracted_path"] = txt_path
        result["extracted_content"] = extracted

        summary_path, summary = summarize_cleaned_text(txt_path)
        result["summary_path"] = summary_path
        result["summary_content"] = summary

        kg_path, kg_content = generate_knowledge_graph(txt_path)
        result["kg_path"] = kg_path
        result["kg_content"] = kg_content

        txt_out, xlsx_out, scripts = generate_test_scripts(kg_path, summary_path)
        result["test_txt_path"] = txt_out
        result["test_xlsx_path"] = xlsx_out
        result["test_scripts"] = scripts

        result["success"] = True

    except Exception as e:
        result["error"] = str(e)

    return result


def get_output_files():
    """Get list of all generated files"""
    files = {
        "extracted": [],
        "summary": [],
        "kg": [],
        "tests_txt": [],
        "tests_xlsx": []
    }
    
    folder_map = {
        "extracted": EXTRACTED_FOLDER,
        "summary": SUMMARY_FOLDER,
        "kg": KG_FOLDER,
        "tests_txt": OUTPUT_TXT_FOLDER,
        "tests_xlsx": OUTPUT_XLSX_FOLDER
    }

    for key, folder in folder_map.items():
        if os.path.exists(folder):
            files[key] = [f for f in os.listdir(folder) if os.path.isfile(os.path.join(folder, f))]

    return files