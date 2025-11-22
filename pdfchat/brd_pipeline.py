"""
BRD -> Summary -> Knowledge Graph -> ABAP Code Pipeline
Converted to importable module for Flask integration
"""

import os
import re
import pdfplumber
from docx import Document
from openai import OpenAI
from langchain.document_loaders import TextLoader
from langchain.text_splitter import TokenTextSplitter
from langchain_openai import ChatOpenAI
from langchain_experimental.graph_transformers import LLMGraphTransformer
from dotenv import load_dotenv

load_dotenv()

api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    raise Exception("ERROR: OPENAI_API_KEY not set in environment!")

client = OpenAI(api_key=api_key)

# Output directories
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(BASE_DIR, "output")
EXTRACTED_FOLDER = os.path.join(OUTPUT_DIR, "extracted")
SUMMARY_FOLDER = os.path.join(OUTPUT_DIR, "summary")
KG_FOLDER = os.path.join(OUTPUT_DIR, "kg")
OUTPUT_TXT_FOLDER = os.path.join(OUTPUT_DIR, "abap_output")

for p in [EXTRACTED_FOLDER, SUMMARY_FOLDER, KG_FOLDER, OUTPUT_TXT_FOLDER]:
    os.makedirs(p, exist_ok=True)

SUMMARY_PROMPT_TEMPLATE = """
You are an expert business analyst specializing in interpreting BRDs, PDDs, FDDs, SOPs and functional documents.

Provide the following:

==========================
PART 1 — FULL DOCUMENT SUMMARY
==========================
Write a clear, complete summary of ALL non-table content. Include:
- Section summaries
- Business purpose
- Key requirements
- Functional logic
- Actors, roles, responsibilities
- System interactions (SAP, VEEVA, S/4HANA, etc.)
- Assumptions and dependencies

==========================
PART 2 — TABLE SUMMARIES (Process Walkthroughs)
==========================
For EACH table, write a 10–12 sentence process explanation covering:
- What the sequence describes
- How the workflow begins and ends
- Role + system interactions
- Decision points
- Validation logic
- Step transitions

DO NOT recreate the table.
Interpret the operational meaning of the steps.

INPUT DOCUMENT — NON-TABLE TEXT:
{clean_text}

INPUT DOCUMENT — TABLES:
{tables_block}
"""


def extract_pdf_content(pdf_path):
    text_content = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                text = page.extract_text()
                if text:
                    text_content += f"\n--- Page {i} Text ---\n{text}\n"
                tables = page.extract_tables()
                for t_i, table in enumerate(tables, start=1):
                    text_content += f"\n--- Table {t_i} ---\n"
                    for row in table:
                        text_content += "\t".join([cell if cell else "" for cell in row]) + "\n"
    except Exception as e:
        raise Exception(f"Error extracting PDF: {e}")
    return text_content


def extract_docx_content(docx_path):
    text_content = ""
    try:
        doc = Document(docx_path)
        for para in doc.paragraphs:
            if para.text and para.text.strip():
                text_content += para.text + "\n"
        for t_i, table in enumerate(doc.tables, start=1):
            text_content += f"\n--- Table {t_i} ---\n"
            for row in table.rows:
                text_content += "\t".join([cell.text.strip() for cell in row.cells]) + "\n"
    except Exception as e:
        raise Exception(f"Error extracting DOCX: {e}")
    return text_content


def process_file_to_txt(file_path):
    base_name, ext = os.path.splitext(os.path.basename(file_path))
    ext = ext.lower()
    
    if ext == ".pdf":
        content = extract_pdf_content(file_path)
    elif ext == ".docx":
        content = extract_docx_content(file_path)
    else:
        raise ValueError("Only PDF/DOCX files supported.")
    
    txt_path = os.path.join(EXTRACTED_FOLDER, f"{base_name}.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write(content)
    return txt_path


def summarize_cleaned_text(txt_path):
    with open(txt_path, "r", encoding="utf-8") as f:
        text = f.read()

    table_pattern = r"---\s*Table\s*(\d+)\s*---(.*?)(?=---\s*Table|\Z)"
    tables = re.findall(table_pattern, text, re.DOTALL)
    clean_text = re.sub(r"---\s*Table\s*\d+.*?(?=---\s*Table|\Z)", "", text, flags=re.DOTALL).strip()

    tables_block = ""
    for table_num, table_content in tables:
        tables_block += f"\n\n### TABLE {table_num}\n{table_content.strip()}"

    prompt = SUMMARY_PROMPT_TEMPLATE.format(clean_text=clean_text, tables_block=tables_block)

    response = client.chat.completions.create(
        model="gpt-4.1",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.0,
    )

    summary = response.choices[0].message.content
    base_name = os.path.splitext(os.path.basename(txt_path))[0]
    summary_path = os.path.join(SUMMARY_FOLDER, f"{base_name}_summary.txt")
    
    with open(summary_path, "w", encoding="utf-8") as f:
        f.write(summary)
    return summary_path, summary


def generate_knowledge_graph(txt_path):
    loader = TextLoader(txt_path, encoding="utf-8")
    document = loader.load()[0]

    text_splitter = TokenTextSplitter(chunk_size=800, chunk_overlap=200)
    chunks = text_splitter.split_documents([document])

    llm = ChatOpenAI(temperature=0, model_name="gpt-4o-mini")
    llm_transformer = LLMGraphTransformer(llm=llm)
    graph_docs = llm_transformer.convert_to_graph_documents(chunks)

    base_name = os.path.splitext(os.path.basename(txt_path))[0]
    out = os.path.join(KG_FOLDER, f"{base_name}_KG.txt")

    total = 0
    kg_content = f"=== KNOWLEDGE GRAPH FOR {base_name} ===\n\n"
    
    for idx, g in enumerate(graph_docs):
        kg_content += f"\n\n######## GRAPH DOCUMENT {idx+1} ########\n"
        kg_content += "\n--- RELATIONSHIPS ---\n"
        for rel in getattr(g, "relationships", []):
            try:
                line = f"{rel.source.id} --{str(rel.type).upper()}--> {rel.target.id}"
                kg_content += line + "\n"
                total += 1
            except:
                pass

    kg_content += f"\nTOTAL RELATIONSHIPS: {total}\n"
    
    with open(out, "w", encoding="utf-8") as f:
        f.write(kg_content)
    return out, kg_content


def generate_abap_code(kg_path, summary_path):
    with open(kg_path, "r", encoding="utf-8") as f:
        graph_text = f.read()
    with open(summary_path, "r", encoding="utf-8") as f:
        summary_text = f.read()

    prompt = f"""
You are an expert SAP ABAP developer.

You will receive TWO inputs:
1. CLEAN SUMMARY extracted from BRD/PDD/FDD
2. CLEAN KNOWLEDGE GRAPH relationships

Using BOTH inputs, generate COMPLETE ABAP CODE that satisfies:

=========================================================
ABAP OUTPUT REQUIREMENTS
=========================================================
- Use realistic SAP S/4HANA objects
- Include:
    * Data declarations
    * SELECT statements
    * Joins based on relationships in the Graph
    * LOOPs to process business logic from Summary
    * VALIDATIONS based on summary requirements
    * FORM or CLASS method implementation
- Use real SAP tables (MARA, VBAK, VBAP, etc.) based on text
- Must be executable-style ABAP logic

STRICT RULES:
- Output ONLY ABAP code — no explanations
- Code should be properly formatted
- No generic placeholders — use real SAP structures

=========================================================
KNOWLEDGE GRAPH:
=========================================================
{graph_text}

=========================================================
SUMMARY DOCUMENT:
=========================================================
{summary_text}

Generate FINAL ABAP CODE now.
"""

    response = client.chat.completions.create(
        model="gpt-4.1",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.1
    )

    abap_code = response.choices[0].message.content
    base_name = os.path.splitext(os.path.basename(kg_path))[0].replace("_KG", "")
    abap_out = os.path.join(OUTPUT_TXT_FOLDER, f"{base_name}_ABAP_Code.txt")

    with open(abap_out, "w", encoding="utf-8") as f:
        f.write(abap_code)
    return abap_out, abap_code


def run_pipeline(file_path):
    """
    Main pipeline function - processes a single BRD/PDD file
    Returns dict with all generated outputs
    """
    result = {
        "success": False,
        "extracted_text_path": None,
        "summary_path": None,
        "summary_content": None,
        "kg_path": None,
        "kg_content": None,
        "abap_path": None,
        "abap_code": None,
        "error": None
    }

    try:
        # Step 1: Extract
        txt_path = process_file_to_txt(file_path)
        result["extracted_text_path"] = txt_path

        # Step 2: Summarize
        summary_path, summary_content = summarize_cleaned_text(txt_path)
        result["summary_path"] = summary_path
        result["summary_content"] = summary_content

        # Step 3: Knowledge Graph
        kg_path, kg_content = generate_knowledge_graph(txt_path)
        result["kg_path"] = kg_path
        result["kg_content"] = kg_content

        # Step 4: ABAP Code
        abap_path, abap_code = generate_abap_code(kg_path, summary_path)
        result["abap_path"] = abap_path
        result["abap_code"] = abap_code

        result["success"] = True

    except Exception as e:
        result["error"] = str(e)

    return result